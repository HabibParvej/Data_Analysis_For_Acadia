import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
import os

# Function to analyze the dataset and generate the report
def generate_report(data_file, output_format="html", selected_columns=None):
    try:
        # Load the dataset
        df = pd.read_csv("sample_dataset_large.csv")
    except Exception as e:
        print(f"Error loading the file: {e}")
        return

    # Create a folder for saving outputs
    output_folder = "report_output"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # File for writing the report
    report_file = os.path.join(output_folder, f"report.{output_format}")
    
    # Analyze dataset
    with open(report_file, "w") as f:
        # List missing values and their percentages
        missing_values = df.isnull().sum()
        missing_percentages = (missing_values / len(df)) * 100
        f.write("Columns with Missing Values:\n")
        f.write(pd.DataFrame({'Missing Count': missing_values, 'Missing %': missing_percentages}).to_string())
        f.write("\n\n")
        print("Columns with Missing Values:")
        print(pd.DataFrame({'Missing Count': missing_values, 'Missing %': missing_percentages}))

        # Categorize columns based on data types
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
        bool_cols = df.select_dtypes(include=['bool']).columns.tolist()
        datetime_cols = df.select_dtypes(include=['datetime']).columns.tolist()

        f.write("Categorized Columns:\n")
        f.write(f"Numeric Columns: {numeric_cols}\n")
        f.write(f"Categorical Columns: {categorical_cols}\n")
        f.write(f"Boolean Columns: {bool_cols}\n")
        f.write(f"Datetime Columns: {datetime_cols}\n")
        f.write("\n")

        print("\nCategorized Columns:")
        print(f"Numeric Columns: {numeric_cols}")
        print(f"Categorical Columns: {categorical_cols}")
        print(f"Boolean Columns: {bool_cols}")
        print(f"Datetime Columns: {datetime_cols}")

        # Remove and list duplicate columns
        duplicate_columns = df.columns[df.T.duplicated()].tolist()
        f.write("Duplicate Columns:\n")
        f.write(f"Before Removing: {list(df.columns)}\n")
        f.write(f"Duplicate Columns Identified: {duplicate_columns}\n")
        df = df.loc[:, ~df.T.duplicated()]
        f.write(f"After Removing: {list(df.columns)}\n\n")
        print("\nDuplicate Columns Removed:")
        print(f"Before: {duplicate_columns}")

        # Remove and list constant columns
        constant_columns = [col for col in df.columns if df[col].nunique() == 1]
        f.write("Constant Columns:\n")
        f.write(f"Before Removing: {list(df.columns)}\n")
        f.write(f"Constant Columns Identified: {constant_columns}\n")
        df = df.drop(columns=constant_columns)
        f.write(f"After Removing: {list(df.columns)}\n\n")
        print("\nConstant Columns Removed:")
        print(f"Before: {constant_columns}")
        print(f"After: {list(df.columns)}")

        # Create boxplots for numeric columns
        boxplot_folder = os.path.join(output_folder, "boxplots")
        if not os.path.exists(boxplot_folder):
            os.makedirs(boxplot_folder)

        for col in numeric_cols:
            if col in df.columns:  # Ensure column exists before plotting
                plt.figure(figsize=(8, 4))
                sns.boxplot(x=df[col])
                plt.title(f"Boxplot for {col}")
                plt.savefig(os.path.join(boxplot_folder, f"{col}_boxplot.png"))
                plt.close()
        print("\nBoxplots saved for numeric columns.")

        # Create distribution charts for specified columns or first 6 columns
        dist_folder = os.path.join(output_folder, "distributions")
        if not os.path.exists(dist_folder):
            os.makedirs(dist_folder)

        columns_to_plot = selected_columns if selected_columns else df.columns[:6]
        for col in columns_to_plot:
            if col in df.columns:  # Ensure column exists before plotting
                plt.figure(figsize=(8, 4))
                if col in numeric_cols:
                    sns.histplot(df[col], kde=True)
                elif col in categorical_cols or col in bool_cols:
                    sns.countplot(x=df[col])
                plt.title(f"Distribution of {col}")
                plt.savefig(os.path.join(dist_folder, f"{col}_distribution.png"))
                plt.close()
        print("\nDistribution charts saved for selected columns.")

    # Embed into PDF report
    if output_format.lower() == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        try:
            with open(report_file, "r") as f:
                for line in f:
                    pdf.multi_cell(0, 10, txt=line)

            # Add boxplots to PDF
            for file in os.listdir(boxplot_folder):
                pdf.add_page()
                pdf.image(os.path.join(boxplot_folder, file), x=10, y=20, w=180)

            # Add distributions to PDF
            for file in os.listdir(dist_folder):
                pdf.add_page()
                pdf.image(os.path.join(dist_folder, file), x=10, y=20, w=180)

            pdf_file = os.path.join(output_folder, "report.pdf")
            pdf.output(pdf_file)
            print(f"PDF report saved at {pdf_file}.")
        except Exception as e:
            print(f"Error while generating PDF: {e}")

    # Generate Word report with images
    elif output_format.lower() == "word":
        doc = Document()
        doc.add_heading('Dataset Analysis Report', level=1)

        # Add text content to Word
        with open(report_file, "r") as f:
            for line in f:
                doc.add_paragraph(line.strip())

        # Add boxplots to Word
        doc.add_heading('Boxplots', level=2)
        for file in os.listdir(boxplot_folder):
            doc.add_picture(os.path.join(boxplot_folder, file), width=doc.sections[0].page_width * 0.7)

        # Add distribution charts to Word
        doc.add_heading('Distributions', level=2)
        for file in os.listdir(dist_folder):
            doc.add_picture(os.path.join(dist_folder, file), width=doc.sections[0].page_width * 0.7)

        word_file = os.path.join(output_folder, "report.docx")
        doc.save(word_file)
        print(f"Word report saved at {word_file}.")

    # Styled HTML output (optional improvement)
    else:
        print(f"HTML report saved at {report_file}.")

# Example Usage
generate_report("sample_dataset.csv", output_format="pdf")
